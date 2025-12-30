"""Utility functions for date parsing, currency detection, and region detection."""
import re
from datetime import datetime, timedelta
from typing import Optional, Dict, Tuple
from io import BytesIO
from docx import Document
from docx.shared import Pt
from models import SalesCallData


# Currency mapping by country
CURRENCY_MAP = {
    "US": "USD", "United States": "USD", "USA": "USD",
    "GB": "GBP", "United Kingdom": "GBP", "UK": "GBP",
    "CA": "CAD", "Canada": "CAD",
    "DE": "EUR", "Germany": "EUR",
    "FR": "EUR", "France": "EUR",
    "IT": "EUR", "Italy": "EUR",
    "ES": "EUR", "Spain": "EUR",
    "NL": "EUR", "Netherlands": "EUR",
    "BE": "EUR", "Belgium": "EUR",
    "AU": "AUD", "Australia": "AUD",
    "JP": "JPY", "Japan": "JPY",
    "CN": "CNY", "China": "CNY",
    "IN": "INR", "India": "INR",
    "BR": "BRL", "Brazil": "BRL",
    "MX": "MXN", "Mexico": "MXN",
    "CH": "CHF", "Switzerland": "CHF",
    "SE": "SEK", "Sweden": "SEK",
    "NO": "NOK", "Norway": "NOK",
    "DK": "DKK", "Denmark": "DKK",
}

# Default currency for common regions
DEFAULT_CURRENCY = "USD"


def parse_quarter_to_date(quarter_text: str, year: Optional[int] = None) -> Optional[str]:
    """
    Parse quarter references to end of quarter date.
    
    Examples:
        "Q2" -> "2025-06-30" (if current year is 2025)
        "Q2 2025" -> "2025-06-30"
        "Q2 of next year" -> "2026-06-30"
        "end of Q3" -> "2025-09-30"
    
    Args:
        quarter_text: Text containing quarter reference
        year: Optional year, if not provided uses current or next year
        
    Returns:
        Date string in YYYY-MM-DD format or None
    """
    if not quarter_text:
        return None
    
    # Extract year from text or use current year
    current_year = datetime.now().year
    if year:
        target_year = year
    else:
        # Check for "next year" or year in text
        if "next year" in quarter_text.lower():
            target_year = current_year + 1
        else:
            year_match = re.search(r'\b(20\d{2})\b', quarter_text)
            if year_match:
                target_year = int(year_match.group(1))
            else:
                target_year = current_year
    
    # Extract quarter number
    quarter_match = re.search(r'Q([1-4])', quarter_text, re.IGNORECASE)
    if not quarter_match:
        return None
    
    quarter = int(quarter_match.group(1))
    
    # Map quarter to end date
    quarter_end_dates = {
        1: (target_year, 3, 31),  # Q1 ends March 31
        2: (target_year, 6, 30),   # Q2 ends June 30
        3: (target_year, 9, 30),   # Q3 ends September 30
        4: (target_year, 12, 31),  # Q4 ends December 31
    }
    
    year, month, day = quarter_end_dates[quarter]
    return f"{year}-{month:02d}-{day:02d}"


def detect_currency_from_location(country: Optional[str] = None, 
                                  city: Optional[str] = None,
                                  region: Optional[str] = None) -> str:
    """
    Detect currency from location information.
    
    Args:
        country: Country name or code
        city: City name (optional, for additional context)
        region: Region name (optional)
        
    Returns:
        Currency code (e.g., "USD", "EUR", "GBP")
    """
    if not country and not region:
        return DEFAULT_CURRENCY
    
    # Check country first
    if country:
        country_upper = country.upper()
        # Check exact match
        if country_upper in CURRENCY_MAP:
            return CURRENCY_MAP[country_upper]
        # Check partial match
        for key, currency in CURRENCY_MAP.items():
            if country_upper in key.upper() or key.upper() in country_upper:
                return currency
    
    # Check region
    if region:
        region_upper = region.upper()
        for key, currency in CURRENCY_MAP.items():
            if region_upper in key.upper() or key.upper() in region_upper:
                return currency
    
    return DEFAULT_CURRENCY


def parse_close_date(date_text: Optional[str]) -> Optional[str]:
    """
    Parse close date from various formats.
    
    Handles:
        - Quarter references (Q2, Q2 2025, etc.)
        - Date strings (YYYY-MM-DD, MM/DD/YYYY, etc.)
        - Relative dates ("in 3 months", "next year", etc.)
    
    Args:
        date_text: Date text from transcript
        
    Returns:
        Date string in YYYY-MM-DD format or None
    """
    if not date_text:
        return None
    
    date_text = date_text.strip()
    
    # Check for quarter references
    if "Q" in date_text.upper():
        return parse_quarter_to_date(date_text)
    
    # Try to parse standard date formats
    date_formats = [
        "%Y-%m-%d",
        "%m/%d/%Y",
        "%d/%m/%Y",
        "%Y/%m/%d",
        "%B %d, %Y",
        "%d %B %Y",
    ]
    
    for fmt in date_formats:
        try:
            dt = datetime.strptime(date_text, fmt)
            return dt.strftime("%Y-%m-%d")
        except ValueError:
            continue
    
    # Handle relative dates
    months_match = re.search(r'(\d+)\s*months?', date_text, re.IGNORECASE)
    if months_match:
        months = int(months_match.group(1))
        future_date = datetime.now() + timedelta(days=months * 30)
        return future_date.strftime("%Y-%m-%d")
    
    return None


def estimate_opportunity_amount(industry: Optional[str] = None,
                               company_size: Optional[int] = None,
                               mentioned_amount: Optional[float] = None) -> Optional[float]:
    """
    Estimate opportunity amount if not explicitly mentioned.
    
    Args:
        industry: Industry type
        company_size: Number of employees
        mentioned_amount: Amount mentioned in transcript
        
    Returns:
        Estimated amount or mentioned amount
    """
    if mentioned_amount:
        return mentioned_amount
    
    # Simple estimation based on company size
    if company_size:
        if company_size < 50:
            return 50000.0  # Small business
        elif company_size < 500:
            return 150000.0  # Mid-market
        elif company_size < 5000:
            return 500000.0  # Enterprise
        else:
            return 1000000.0  # Large enterprise
    
    return None


def format_opportunity_name(account_name: str, project_name: Optional[str] = None) -> str:
    """
    Format opportunity name as "Customer name - project name".
    
    Args:
        account_name: Account/company name
        project_name: Project or offering name
        
    Returns:
        Formatted opportunity name
    """
    if project_name:
        return f"{account_name} - {project_name}"
    return account_name


def detect_lead_source(transcript: str) -> str:
    """
    Detect lead source from transcript.
    
    Args:
        transcript: Call transcript
        
    Returns:
        Lead source (defaults to "Call", "partner-google" if Google mentioned)
    """
    transcript_lower = transcript.lower()
    
    # Check if Google/Googler mentioned
    if any(keyword in transcript_lower for keyword in ["google", "googler", "gcp"]):
        return "partner-google"
    
    return "Call"


def format_deal_summary(summary: str, next_steps: Optional[str] = None) -> str:
    """
    Format deal summary with timestamp.
    
    Args:
        summary: Call summary
        next_steps: Next steps (optional)
        
    Returns:
        Formatted deal summary with timestamp
    """
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    formatted = f"[{timestamp}] Initial call summary: {summary}"
    
    if next_steps:
        formatted += f"\n\nNext Steps: {next_steps}"
    
    return formatted


def generate_deal_summary_docx(call_data: SalesCallData) -> BytesIO:
    """
    Generate a Word document with deal summary following the structure from summary.docx.
    
    Args:
        call_data: SalesCallData containing account, contact, opportunity, and MEDDIC information
        
    Returns:
        BytesIO object containing the docx file
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Deal Summary', 0)
    
    # Section 1: Executive Summary (The Punchline)
    doc.add_heading('1. Executive Summary (The Punchline)', level=2)
    
    # Key Synthesis (2-3 Sentences)
    doc.add_heading('Key Synthesis (2-3 Sentences)', level=3)
    if call_data.opportunity and call_data.opportunity.deal_summary:
        # Extract first 2-3 sentences from deal_summary
        summary_text = call_data.opportunity.deal_summary
        # Remove timestamps if present
        summary_text = re.sub(r'\[.*?\]\s*', '', summary_text)
        sentences = re.split(r'[.!?]+', summary_text)
        sentences = [s.strip() for s in sentences if s.strip()]
        key_synthesis = '. '.join(sentences[:3])
        if key_synthesis and not key_synthesis.endswith('.'):
            key_synthesis += '.'
        doc.add_paragraph(key_synthesis if key_synthesis else "To be completed based on call transcript.")
    elif call_data.summary:
        doc.add_paragraph(call_data.summary)
    else:
        doc.add_paragraph("To be completed based on call transcript.")
    
    # Section 2: Situation & Client Context (Why Now?)
    doc.add_heading('2. Situation & Client Context (Why Now?)', level=2)
    
    # Client Overview
    client_name = call_data.account.name if call_data.account and call_data.account.name else "[Insert Client/Area Focus]"
    doc.add_heading(f'Client Overview: {client_name}', level=3)
    client_overview_parts = []
    if call_data.account:
        account = call_data.account
        if account.name:
            client_overview_parts.append(f"Client: {account.name}")
        if account.industry:
            client_overview_parts.append(f"Industry: {account.industry}")
        if account.billing_city and account.billing_country:
            client_overview_parts.append(f"Location: {account.billing_city}, {account.billing_country}")
        if account.annual_revenue:
            client_overview_parts.append(f"Annual Revenue: ${account.annual_revenue:,.0f}")
        if account.number_of_employees:
            client_overview_parts.append(f"Employees: {account.number_of_employees:,}")
    if call_data.contact:
        contact = call_data.contact
        if contact.first_name or contact.last_name:
            name_parts = [p for p in [contact.first_name, contact.last_name] if p]
            client_overview_parts.append(f"Primary Contact: {' '.join(name_parts)}")
        if contact.title:
            client_overview_parts.append(f"Title: {contact.title}")
    doc.add_paragraph('\n'.join(client_overview_parts) if client_overview_parts else "To be completed.")
    
    # Market Position & Core Challenge
    doc.add_heading('Market Position & Core Challenge (including Revenue)', level=3)
    challenge_parts = []
    if call_data.opportunity and call_data.opportunity.meddic:
        meddic = call_data.opportunity.meddic
        if meddic.identified_pain:
            challenge_parts.append(meddic.identified_pain)
    if call_data.account and call_data.account.annual_revenue:
        challenge_parts.append(f"Company Revenue: ${call_data.account.annual_revenue:,.0f}")
    doc.add_paragraph('\n\n'.join(challenge_parts) if challenge_parts else "To be completed based on call discussion.")
    
    # Strategic/Technology Context
    doc.add_heading('Strategic/Technology Context', level=3)
    if call_data.opportunity and call_data.opportunity.meddic:
        meddic = call_data.opportunity.meddic
        context_parts = []
        if meddic.decision_criteria_notes:
            context_parts.append(f"Decision Criteria: {meddic.decision_criteria_notes}")
        if meddic.decision_process_notes:
            context_parts.append(f"Decision Process: {meddic.decision_process_notes}")
        doc.add_paragraph('\n\n'.join(context_parts) if context_parts else "To be completed.")
    else:
        doc.add_paragraph("To be completed.")
    
    # Section 3: The Project & Immediate Value (The Solution)
    doc.add_heading('3. The Project & Immediate Value (The Solution)', level=2)
    
    # Project Objective
    doc.add_heading('Project Objective', level=3)
    if call_data.opportunity:
        opp = call_data.opportunity
        if opp.name:
            # Extract project name from opportunity name (format: "Customer - Project")
            if ' - ' in opp.name:
                project_name = opp.name.split(' - ', 1)[1]
                doc.add_paragraph(project_name)
            else:
                doc.add_paragraph(opp.name)
        else:
            doc.add_paragraph("To be completed.")
    else:
        doc.add_paragraph("To be completed.")
    
    # Project Name MVP
    if call_data.opportunity and call_data.opportunity.name:
        if ' - ' in call_data.opportunity.name:
            project_name = call_data.opportunity.name.split(' - ', 1)[1]
        else:
            project_name = call_data.opportunity.name
        doc.add_heading(f'{project_name} MVP', level=3)
        doc.add_paragraph(f"{project_name} MVP")
    else:
        doc.add_heading('[Insert Project Name] MVP', level=3)
        doc.add_paragraph("To be completed.")
    
    # Scope and Timeline
    doc.add_heading('Scope and Timeline', level=4)
    scope_parts = []
    if call_data.opportunity:
        opp = call_data.opportunity
        if opp.projected_start_date:
            scope_parts.append(f"Start Date: {opp.projected_start_date}")
        if opp.close_date:
            scope_parts.append(f"Close Date: {opp.close_date}")
        if opp.number_of_weeks:
            scope_parts.append(f"Duration: {opp.number_of_weeks} weeks")
        if opp.amount:
            scope_parts.append(f"Project Value: ${opp.amount:,.0f}")
    doc.add_paragraph('\n'.join(scope_parts) if scope_parts else "To be completed.")
    
    # Resources needed
    doc.add_heading('Resources needed (if discussed)', level=4)
    doc.add_paragraph("To be completed based on call discussion.")
    
    # Immediate Business Value & Metrics
    doc.add_heading('Immediate Business Value & Metrics', level=3)
    value_parts = []
    if call_data.opportunity and call_data.opportunity.meddic:
        meddic = call_data.opportunity.meddic
        if meddic.metrics_notes:
            value_parts.append(meddic.metrics_notes)
    if call_data.opportunity and call_data.opportunity.next_steps:
        value_parts.append(f"Next Steps: {call_data.opportunity.next_steps}")
    doc.add_paragraph('\n\n'.join(value_parts) if value_parts else "To be completed.")
    
    # Section 4: Strategic Impact & Path to Scale (The Future)
    doc.add_heading('4. Strategic Impact & Path to Scale (The Future)', level=2)
    
    # Strategic Value and Expansion Potential
    doc.add_heading('Strategic Value and Expansion Potential', level=3)
    if call_data.opportunity and call_data.opportunity.meddic:
        meddic = call_data.opportunity.meddic
        if meddic.economic_buyers_notes:
            doc.add_paragraph(f"Economic Buyer: {meddic.economic_buyers_notes}")
        if meddic.champion:
            doc.add_paragraph(f"Champion: {meddic.champion}")
    doc.add_paragraph("Strategic value and expansion opportunities to be discussed.")
    
    # Scaling Roadmap
    doc.add_heading('Scaling Roadmap (Future Phases)', level=4)
    doc.add_paragraph("To be completed based on strategic discussions.")
    
    # Adjacent Use Case Opportunities
    doc.add_heading('Adjacent Use Case Opportunities', level=4)
    doc.add_paragraph("To be identified and documented.")
    
    # Save to BytesIO
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes
